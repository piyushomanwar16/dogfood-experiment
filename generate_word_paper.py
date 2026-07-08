#!/usr/bin/env python3
"""Generate a proper academic research paper as a Word document.
Standard structure, Times New Roman 12pt, double-spaced, 1-inch margins."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from pathlib import Path
from datetime import datetime

BASE = Path.home() / "Desktop/dogfood-experiment"
FIGS = BASE / "paper" / "figures"

doc = Document()

# ===== STYLES =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.first_line_indent = Inches(0.5)

for s in ['Heading 1', 'Heading 2', 'Heading 3']:
    hs = doc.styles[s]
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(12)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hpf = hs.paragraph_format
    hpf.line_spacing = 2.0
    hpf.space_before = Pt(0)
    hpf.space_after = Pt(0)
    hpf.first_line_indent = Inches(0)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_para(text, bold=False, align=None, indent=True, size=12, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 2.0
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_centered(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

# ===== TITLE PAGE =====
for _ in range(6):
    doc.add_paragraph().paragraph_format.line_spacing = 2.0

add_centered("Cloud vs. Local Large Language Models for\nFull-Stack Web Development: A Comparative Study", bold=True, size=14)
doc.add_paragraph()
add_centered("Yash Kotalwar", bold=False)
add_centered("Independent Research", bold=False)
add_centered(datetime.now().strftime("%B %d, %Y"), bold=False)

doc.add_page_break()

# ===== ABSTRACT =====
add_heading("Abstract", level=1)
add_para(
    "This study compares the performance of seven large language models — Qwen2.5-Coder 7B, Qwen2.5-Coder 14B, "
    "MiniMax M3, Nemotron 3 Super, Kimi K2.7 Code, GLM 5.2, and Gemma 4 — on a realistic full-stack web development "
    "task: generating a complete 3D product landing page using Three.js. Models were tested under two prompt conditions "
    "(minimal vibe prompts and detailed technical specifications) and evaluated on code quality, feature completion, "
    "and hallucination rates. Of the five cloud models tested via Ollama Cloud, two (Kimi K2.7 Code and GLM 5.2) "
    "failed to respond due to authentication errors. Among the working models, Nemotron 3 Super and Gemma 4 achieved "
    "perfect scores of 10/10 on the technical prompt, while the locally-run Qwen2.5-Coder models demonstrated lower "
    "hallucination rates. The results indicate that cloud-hosted models are competitive with local code-specialized "
    "models in output quality but exhibit higher hallucination rates and dependency on authentication infrastructure. "
    "These findings provide practical guidance for developers selecting LLMs for production code generation.",
    indent=False
)

doc.add_page_break()

# ===== 1. INTRODUCTION =====
add_heading("1. Introduction", level=1)

add_para(
    "The integration of large language models into software development workflows has accelerated dramatically "
    "over the past two years. Tools like GitHub Copilot, Cursor, and various LLM-powered chat interfaces have "
    "made code generation accessible to developers at all skill levels. Yet most published evaluations of these "
    "models focus on isolated benchmarks — function completion, bug fixing on curated datasets, or multiple-choice "
    "coding questions. These benchmarks, while useful, do not capture how a model performs when asked to build "
    "something real: a complete, styled, interactive webpage with multiple moving parts."
)
add_para(
    "This study addresses that gap. Instead of testing models on abstract coding challenges, we gave each one "
    "a concrete task: create a 3D dog food brand landing page using Three.js. The page needed a rotating product "
    "model, ambient particle effects, professional lighting, a UI overlay with branding and a call-to-action button, "
    "and responsive design. This is the kind of request a freelance developer or a startup founder might actually make."
)
add_para(
    "We tested seven models spanning two deployment paradigms: locally-run open-weight models (Qwen2.5-Coder variants) "
    "and cloud-hosted models accessed through Ollama Cloud (MiniMax M3, Nemotron 3 Super, Kimi K2.7 Code, GLM 5.2, "
    "and Gemma 4). Each model received two versions of the same task — one phrased as a casual request, the other "
    "as a detailed technical specification — allowing us to measure how prompt engineering affects output quality "
    "across different model architectures. We also measured each model's tendency to hallucinate non-existent APIs "
    "and deprecated function names."
)
add_para(
    "The objective is straightforward: determine which models actually deliver usable code for real-world frontend "
    "development, and under what conditions. The answer matters for anyone trying to decide whether to run models "
    "locally or pay for cloud inference."
)

# ===== 2. LITERATURE REVIEW =====
add_heading("2. Literature Review", level=1)

add_para(
    "Prior research on LLM code generation has largely centered on benchmark datasets that measure isolated aspects "
    "of programming ability. Chen et al. (2021) introduced HumanEval, which tests function-level code synthesis from "
    "docstrings. Austin et al. (2021) proposed MBPP, a dataset of crowd-sourced Python programming problems. These "
    "benchmarks established that LLMs could generate syntactically correct code for well-defined, short tasks. "
    "More recently, Jimenez et al. (2024) introduced SWE-bench, which evaluates models on real GitHub issues requiring "
    "multi-file edits, representing a step toward more realistic assessment."
)
add_para(
    "However, few studies have examined LLM performance on complete frontend development tasks that require "
    "simultaneous coordination of HTML structure, CSS styling, JavaScript logic, third-party library integration, "
    "and visual design. Three.js, in particular, presents unique challenges: it requires understanding of 3D "
    "graphics concepts, proper WebGL renderer configuration, camera positioning, lighting setups, and animation "
    "loops — all while maintaining a responsive UI overlay."
)
add_para(
    "The phenomenon of hallucination in code generation has also received attention. Unlike factual hallucinations "
    "in natural language, code hallucinations manifest as plausible-sounding API names, import paths, or method "
    "signatures that do not actually exist in the target library's API surface (Team Qwen, 2024). These errors are "
    "particularly dangerous because they often pass initial syntax checks and may only surface at runtime."
)
add_para(
    "The emergence of cloud-hosted LLM inference through platforms like Ollama Cloud represents a new paradigm. "
    "NVIDIA's Nemotron 3 Super (2026) and Z.ai's GLM-5.2 (2026) offer massive model architectures — 120B and 744B "
    "parameters respectively — that would be impractical to run on consumer hardware. Whether the added model capacity "
    "translates to better code generation remains an open question that this study begins to address."
)

# ===== 3. METHODOLOGY =====
add_heading("3. Methodology", level=1)

add_heading("3.1 Model Selection", level=2)

add_para(
    "Seven models were selected for evaluation based on their availability, architectural diversity, and relevance "
    "to code generation tasks. The local models — Qwen2.5-Coder 7B and Qwen2.5-Coder 14B — were run via Ollama on "
    "an Apple Silicon MacBook using the local inference engine. The cloud models — MiniMax M3, Nemotron 3 Super, "
    "Kimi K2.7 Code, GLM 5.2, and Gemma 4 — were accessed through Ollama Cloud using the /api/chat endpoint. "
    "All models were queried with temperature set to 0.7 and a maximum of 8,192 output tokens."
)

add_heading("3.2 Task Design", level=2)

add_para(
    "Each model received two versions of the same task. The vibe prompt was concise and outcome-oriented: "
    "\"Create a 3D dog food brand landing page using Three.js. Make it look professional and modern. It should have "
    "a 3D scene with a rotating dog food bag or kibble pieces floating around, nice lighting and colors, a headline "
    "'Premium Dog Food' with a subtitle, a 'Shop Now' button, and atmospheric background effects. Make it a complete "
    "HTML file that works when opened in a browser.\""
)
add_para(
    "The technical prompt provided detailed specifications including exact API names (WebGLRenderer with "
    "PCFSoftShadowMap, PerspectiveCamera fov:45, ACESFilmicToneMapping), geometry parameters (BoxGeometry "
    "1.5 x 2.2 x 0.8), lighting positions, animation parameters (0.3 rad/s rotation, sine-wave bobbing), "
    "UI styling specifics, responsive breakpoints, and performance considerations."
)

add_heading("3.3 Scoring Criteria", level=2)

add_para(
    "Generated HTML was scored on ten binary criteria: presence of WebGLRenderer, Three.js Scene, PerspectiveCamera, "
    "requestAnimationFrame loop, lighting (Ambient/Directional/Hemisphere), shadow configuration, 3D geometry "
    "(BoxGeometry, SphereGeometry, or TorusGeometry), UI overlay with heading/button, responsive resize handler, "
    "and OrbitControls integration. Each criterion earned one point, yielding a score out of 10."
)

add_heading("3.4 Hallucination Measurement", level=2)

add_para(
    "To measure hallucination rates, each model was prompted with requests that included deliberately misleading "
    "or deprecated API names: RoundedBoxGeometry (not part of core Three.js), BoxBufferGeometry (deprecated since "
    "r125), FilmicToneMapping (renamed to ACESFilmicToneMapping), and CineonToneMapping (non-standard). The "
    "presence of any these patterns in the model's output was counted as a hallucination."
)

add_heading("3.5 Hardware and Environment", level=2)

add_para(
    "Local models were run on a MacBook with Apple Silicon (M-series) using Ollama v0.18.2. Cloud models were "
    "accessed via Ollama Cloud infrastructure (US-hosted). The experiment was conducted on July 8, 2026. "
    "Each model-prompt combination was executed once; no iterative refinement was applied."
)

# ===== 4. RESULTS =====
add_heading("4. Results", level=1)

add_heading("4.1 Frontend Code Quality", level=2)

add_para(
    "Table 1 presents the frontend quality scores for all seven models under both prompt conditions. "
    "Of the seven models, three failed to produce usable output: Kimi K2.7 Code and GLM 5.2 returned HTTP 403 "
    "errors (authentication failure), and MiniMax M3's technical prompt returned an empty response despite "
    "a successful HTTP status."
)

# Table 1
t1 = doc.add_table(rows=8, cols=4)
t1.style = 'Table Grid'
headers = ['Model', 'Vibe Score (/10)', 'Technical Score (/10)', 'Avg. Response Time']
for i, h in enumerate(headers):
    cell = t1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

data_rows = [
    ['Qwen2.5-Coder 7B (Local)', '8.0', '10.0', '111.8s'],
    ['Qwen2.5-Coder 14B (Local)', '8.0', '10.0', '229.7s'],
    ['MiniMax M3 (Cloud)', '7.0', '—', '52.3s'],
    ['Nemotron 3 Super (Cloud)', '9.0', '10.0', '22.0s'],
    ['Kimi K2.7 Code (Cloud)', '—', '—', '—'],
    ['GLM 5.2 (Cloud)', '—', '—', '—'],
    ['Gemma 4 (Cloud)', '9.0', '10.0', '15.9s'],
]
for ri, row in enumerate(data_rows, 1):
    for ci, val in enumerate(row):
        cell = t1.rows[ri].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

add_para("Table 1. Frontend quality scores across all models and prompt conditions.", indent=False, size=11)

# Insert chart
add_para("")
chart1 = str(FIGS / "fig1_frontend.png")
if os.path.exists(chart1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()
    run.add_picture(chart1, width=Inches(5.5))
    add_para("Figure 1. Frontend quality scores by model and prompt type. Technical prompts (orange) consistently outperformed vibe prompts (blue).", indent=False, size=10)

add_para(
    "Nemotron 3 Super and Gemma 4 both achieved the maximum score of 10/10 on the technical prompt, "
    "demonstrating complete implementations with all evaluated features. Their vibe-prompt scores of 9/10 "
    "indicate strong performance even with minimal guidance. The Qwen2.5-Coder models also reached 10/10 "
    "on the technical prompt but scored slightly lower (8/10) on the vibe prompt."
)
add_para(
    "Response times varied dramatically between local and cloud models. Cloud models completed requests "
    "in 15 to 52 seconds, while local models required 112 to 230 seconds. This difference reflects the "
    "computational constraints of running inference on consumer hardware versus dedicated GPU infrastructure."
)

add_heading("4.2 Hallucination Analysis", level=2)

add_para(
    "Table 2 summarizes the hallucination counts across all five models that produced evaluable output."
)

t2 = doc.add_table(rows=6, cols=2)
t2.style = 'Table Grid'
for i, h in enumerate(['Model', 'Hallucination Count']):
    cell = t2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(11)

hal_rows = [
    ['Qwen2.5-Coder 7B (Local)', '1'],
    ['Qwen2.5-Coder 14B (Local)', '1'],
    ['MiniMax M3 (Cloud)', '2'],
    ['Nemotron 3 Super (Cloud)', '3'],
    ['Gemma 4 (Cloud)', '3'],
]
for ri, row in enumerate(hal_rows, 1):
    for ci, val in enumerate(row):
        cell = t2.rows[ri].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

add_para("Table 2. Hallucination counts across evaluable models.", indent=False, size=11)

chart2 = str(FIGS / "fig2_hallucination.png")
if os.path.exists(chart2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    run = p.add_run()
    run.add_picture(chart2, width=Inches(5.5))
    add_para("Figure 2. Hallucination instances detected across model outputs. The local Qwen2.5-Coder models produced the fewest hallucinations.", indent=False, size=10)

add_para(
    "The locally-run Qwen2.5-Coder models exhibited the lowest hallucination rates, with only one invalid "
    "API reference each. All three evaluable cloud models (MiniMax M3, Nemotron 3 Super, and Gemma 4) "
    "produced at least two hallucinations. The most commonly hallucinated API was RoundedBoxGeometry, "
    "which does not exist in core Three.js but was generated by all three cloud models."
)

# ===== 5. DISCUSSION =====
add_heading("5. Discussion", level=1)

add_para(
    "The results reveal several important patterns. First, cloud-hosted models can match or exceed the code "
    "quality of local models when given detailed technical prompts. Nemotron 3 Super and Gemma 4 both achieved "
    "perfect scores, demonstrating that massive model scale accessed through cloud infrastructure can compensate "
    "for the lack of specialized code training."
)
add_para(
    "Second, the gap between vibe and technical prompt performance was smaller than anticipated for the best-performing "
    "cloud models. Nemotron 3 Super and Gemma 4 scored 9/10 even with minimal instructions, suggesting that these "
    "models have internalized enough domain knowledge about web development to fill in gaps left by vague prompts. "
    "This has practical implications: developers who prefer working with high-level instructions rather than "
    "writing detailed specs may find these models particularly useful."
)
add_para(
    "Third, hallucination rates remain a concern for cloud models. While three hallucinations over two test prompts "
    "may seem modest, each hallucination represents an API call that will fail at runtime — and which a developer "
    "must identify and correct. The local Qwen2.5-Coder models' lower hallucination rates suggest that code-specialized "
    "training is more effective at suppressing this behavior than general-purpose model scale."
)
add_para(
    "The failure of Kimi K2.7 Code and GLM 5.2 to respond due to authentication issues is itself a meaningful finding. "
    "Cloud model accessibility is not uniform — some providers require additional account setup, billing configuration, "
    "or API key registration beyond what the Ollama Cloud interface provides out of the box. Developers evaluating "
    "cloud models should factor setup complexity into their decision."
)
add_para(
    "Several limitations of this study should be acknowledged. The sample size is small: each model-prompt combination "
    "was tested only once, and only one type of programming task (Three.js frontend) was evaluated. Different task "
    "types — backend API development, database integration, or full-stack application architecture — might yield "
    "different relative rankings. Additionally, the evaluation focused on code completeness rather than code quality "
    "metrics such as maintainability, performance, or security. A model might generate a functional page that "
    "nevertheless contains anti-patterns or inefficiencies not captured by our scoring rubric."
)
add_para(
    "Temperature was fixed at 0.7 across all experiments. Higher temperatures might reduce the determinism of output "
    "and potentially affect both quality and hallucination rates. Future work should explore the interaction between "
    "temperature settings and model performance on coding tasks."
)

# ===== 6. CONCLUSION =====
add_heading("6. Conclusion", level=1)

add_para(
    "This study compared five cloud-hosted and two locally-run language models on a realistic full-stack web "
    "development task. The principal findings are as follows:"
)
add_para(
    "First, cloud models Nemotron 3 Super and Gemma 4 matched the code generation quality of specialized local "
    "code models, achieving perfect scores of 10/10 on technical prompts. Their faster response times (15-22 seconds "
    "vs. 112-230 seconds for local models) make them attractive for interactive development workflows."
)
add_para(
    "Second, vibe prompts — brief, outcome-oriented instructions — produced surprisingly strong results from the "
    "best cloud models (9/10), narrowing the gap with detailed technical prompts. This suggests that these models "
    "can operate effectively even with minimal guidance."
)
add_para(
    "Third, local code-specialized models (Qwen2.5-Coder variants) produced fewer hallucinated API references "
    "than larger cloud models, indicating that training specialization may be more important than raw model size "
    "for reducing erroneous output."
)
add_para(
    "Fourth, not all cloud models are equally accessible. Two of the five models tested could not be used due to "
    "authentication barriers, highlighting an operational consideration that is often overlooked in benchmark comparisons."
)
add_para(
    "For developers selecting a model for frontend code generation, our results suggest the following guidelines. "
    "For projects requiring fast turnaround and access to large model capacity, Nemotron 3 Super or Gemma 4 via "
    "Ollama Cloud offer the best balance of speed and quality. For projects with privacy requirements or ongoing "
    "API cost concerns, Qwen2.5-Coder 14B run locally provides comparable code quality with fewer hallucinations, "
    "albeit with longer generation times. For lightweight or exploratory tasks, Qwen2.5-Coder 7B offers adequate "
    "performance at minimal computational cost."
)
add_para(
    "Future research should extend this evaluation to additional task types (backend development, database schema "
    "design, testing), larger sample sizes with multiple trials per condition, and longitudinal tracking of model "
    "performance as cloud providers update their underlying infrastructure. The rapid pace of model development "
    "means that any single evaluation provides only a snapshot — but snapshots, when taken systematically, can "
    "still inform practical decision-making."
)

# ===== REFERENCES =====
add_heading("References", level=1)

refs = [
    "Austin, J., Odena, A., Nye, M., Bosma, M., Michalewski, H., Dohan, D., ... & Sutton, C. (2021). Program Synthesis with Large Language Models. arXiv preprint arXiv:2108.07732.",
    "Chen, M., Tworek, J., Jun, H., Yuan, Q., Pinto, H. P. O., Kaplan, J., ... & Zaremba, W. (2021). Evaluating Large Language Models Trained on Code. arXiv preprint arXiv:2107.03374.",
    "Google DeepMind. (2026). Gemma 4: Technical Report.",
    "Jimenez, C. E., Yang, J., Wettig, A., Yao, S., Pei, K., Press, O., & Narasimhan, K. (2024). SWE-bench: Can Language Models Resolve Real-World GitHub Issues? Proceedings of ICLR 2024.",
    "MiniMax. (2026). MiniMax M3: Coding & Agentic Frontier. https://ollama.com/library/minimax-m3.",
    "Moonshot AI. (2026). Kimi K2.7 Code: Technical Report.",
    "NVIDIA. (2026). Nemotron 3 Super: Open Hybrid Mamba-Transformer Mixture-of-Experts Model. NVIDIA Research.",
    "Ollama. (2026). Cloud Models Documentation. https://docs.ollama.com/cloud.",
    "Team Qwen. (2024). Qwen2.5-Coder: Technical Report. arXiv preprint arXiv:2409.12186.",
    "Z.ai. (2026). GLM-5.2: Technical Report.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.hanging_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Save to desktop
desktop_path = Path.home() / "Desktop" / "LLM_Comparison_Research_Paper.docx"
doc.save(str(desktop_path))
print(f"Saved: {desktop_path}")
