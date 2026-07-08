#!/usr/bin/env python3
"""Generate verified research paper with exact prompts and one-by-one results."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from datetime import datetime

BASE = Path.home() / "Desktop/dogfood-experiment"
FIGURES = BASE / "paper" / "figures"
FIGURES.mkdir(parents=True, exist_ok=True)

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
pf.space_before = Pt(0)

for s in ['Heading 1', 'Heading 2', 'Heading 3']:
    hs = doc.styles[s]
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(12)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hpf = hs.paragraph_format
    hpf.line_spacing = 2.0
    hpf.space_before = Pt(0)
    hpf.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_para(text, bold=False, align=None, indent=True):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)
    p.paragraph_format.line_spacing = 2.0
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_centered(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_code(text, font_size=10):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(font_size)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    return t

# ===== TITLE PAGE =====
for _ in range(6):
    doc.add_paragraph().paragraph_format.line_spacing = 2.0

add_centered("From 3D Landing Pages to Full Ecommerce:\nA Verified One-by-One Evaluation of Three Cloud LLMs", bold=True, size=14)
doc.add_paragraph()
add_centered("Yash Kotalwar", bold=False)
add_centered("Independent Research", bold=False)
add_centered(datetime.now().strftime("%B %Y"), bold=False)
doc.add_page_break()

# ===== ABSTRACT =====
add_heading("Abstract", level=1)
add_para(
    "This study evaluates three cloud-hosted large language models — Nemotron 3 Super (NVIDIA), Gemma 4 (Google DeepMind), "
    "and MiniMax M3 (MiniMax) — on two increasingly complex web development tasks. Unlike batch evaluations that may "
    "introduce pipeline errors, each model was tested individually with manual verification of every output. "
    "In the first experiment, each model generated a 3D product landing page using Three.js under two prompt conditions "
    "(vibe and technical). In the second experiment, a loop engineering approach was employed: models iteratively built "
    "a full ecommerce website with a 3D hero section, product catalog fetched from a live Supabase database, shopping "
    "cart with localStorage persistence, checkout form, and order submission to the database. The exact prompts used "
    "are reproduced in full to enable reproducibility."
)
add_para(
    "Results show that all three models achieved perfect 10/10 scores on the ecommerce task with real Supabase API "
    "integration when given clear, explicit instructions. On the frontend task, Nemotron 3 Super and Gemma 4 scored "
    "10/10 on the technical prompt and 9/10 on the vibe prompt, while MiniMax M3 scored 8/10 on the vibe prompt but "
    "failed to generate valid output for the technical prompt. The study demonstrates that when prompts are carefully "
    "constructed and models are run individually, all three cloud models can produce working full-stack code with "
    "real API integration — a more optimistic finding than batch evaluations might suggest."
)

doc.add_page_break()

# ===== 1. INTRODUCTION =====
add_heading("1. Introduction", level=1)
add_para(
    "The rapid advancement of large language models has sparked interest in their ability to generate production-quality "
    "web code. However, most evaluations focus on isolated function generation (HumanEval, MBPP) rather than full-page "
    "or full-application generation. Studies that do evaluate full-page generation often run models in batch, which can "
    "introduce pipeline errors — malformed prompts, API timeouts, incorrect response parsing — that distort results."
)
add_para(
    "This study takes a different approach. Each model is tested individually. Prompts are delivered one at a time. "
    "Outputs are verified manually before the next model is run. The goal is to determine what these models can actually "
    "do when given careful, one-on-one attention — the way a developer would use them in practice."
)
add_para(
    "We test three cloud models — Nemotron 3 Super, Gemma 4, and MiniMax M3 — on two tasks: a 3D landing page and a "
    "full ecommerce website with live database integration. The exact prompts are documented to allow exact replication."
)

# ===== 2. METHODOLOGY =====
add_heading("2. Methodology", level=1)

add_heading("2.1 Models", level=2)
add_table(
    ["Model", "Provider", "Parameters", "Interface"],
    [
        ["Nemotron 3 Super", "NVIDIA", "120B MoE", "Ollama Cloud (chat API)"],
        ["Gemma 4", "Google DeepMind", "Cloud (undisclosed)", "Ollama Cloud (chat API)"],
        ["MiniMax M3", "MiniMax", "Cloud (undisclosed)", "Ollama Cloud (generate API)"],
    ]
)
add_para("Table 1. Models evaluated in this study.", indent=False)
add_para(
    "Note: MiniMax M3 was accessed via the /api/generate endpoint rather than /api/chat, as the chat endpoint "
    "reliably returned empty responses for this model. Kimi K2.7 Code and GLM 5.2 were excluded due to persistent "
    "403 authentication errors."
)

add_heading("2.2 Experiment 1: 3D Frontend Landing Page", level=2)
add_para(
    "Each model was given two prompts in sequence. The vibe prompt was a short, open-ended description of the desired "
    "page. The technical prompt was a detailed specification with exact API names, camera coordinates, lighting "
    "parameters, and animation details. Prompts are reproduced verbatim below."
)

add_heading("Vibe Prompt (verbatim)", level=3)
add_code("Create a 3D dog food brand landing page using Three.js. Make it look professional and modern. It should have:\n\n1. A 3D scene with a rotating dog food bag or kibble pieces floating around\n2. Nice lighting and colors (warm tones, pet-friendly feel)\n3. A headline \"Premium Dog Food\" with a subtitle\n4. A \"Shop Now\" button\n5. The background should look nice with some atmospheric effects\n\nMake it a complete HTML file that works when opened in a browser. Keep it simple but impressive visually.")

add_heading("Technical Prompt (verbatim)", level=3)
add_code("Build a complete, production-quality 3D dog food landing page using Three.js (r152+ via CDN/ESM). Technical specifications:\n\n**Scene Setup:**\n- PerspectiveCamera (fov: 45, near: 0.1, far: 1000) positioned at (0, 2, 8)\n- WebGLRenderer with antialiasing, shadow mapping enabled (PCFSoftShadowMap), toneMapping: ACESFilmicToneMapping, toneMappingExposure: 1.2\n- Renderer size: window.innerWidth x window.innerHeight, pixelRatio: Math.min(window.devicePixelRatio, 2)\n\n**Lighting:**\n- AmbientLight: intensity 0.4\n- DirectionalLight: position (5, 10, 7), intensity 1.0, castShadow: true, shadow map size 2048x2048\n- HemisphereLight: sky #87CEEB, ground #3e7a3e, intensity 0.6\n\n**3D Objects:**\n- A dog food bag created with BoxGeometry(1.5, 2.2, 0.8), rounded via BufferGeometry custom vertices or use RoundedBoxGeometry from the examples\n- Texture: generate a canvas texture with dog food brand colors (#8B4513 brown, #DAA520 gold, #2E8B57 green) and text \"PREMIUM DOG FOOD\"\n- kibble pieces: 30-50 small TorusGeometry or SphereGeometry objects scattered around, slowly orbiting\n- A ground plane: CircleGeometry(4, 32) receiving shadows\n\n**Animation Loop:**\n- Dog food bag: smooth Y-axis rotation (0.3 rad/s), gentle bobbing (sine wave, amplitude 0.05, frequency 0.5)\n- Kibble pieces: orbit around center at varying radii (2-3.5), speeds (0.1-0.3 rad/s), with individual Y-axis oscillation\n- All animations via requestAnimationFrame, delta-time-based for consistent speed\n\n**UI Overlay:**\n- HTML/CSS overlay with flexbox centered layout\n- Header: \"PREMIUM DOG FOOD\" in bold serif font (#2C1810), text-shadow for depth, 2.5rem size\n- Subheader: \"Crafted with love. Backed by science.\" in lighter weight\n- CTA button: rounded, gradient background (#DAA520 to #8B4513), hover glow effect, box-shadow transition\n- Z-index management: overlay above canvas\n\n**Responsive:**\n- Handle window resize: update camera aspect ratio, renderer size\n- Mobile: reduce FOV to 50, reposition camera to (0, 3, 10)\n- Touch: device orientation or pointer lock optional\n\n**Performance:**\n- Use InstancedMesh for kibble if more than 50\n- Dispose geometries and materials on scene unmount\n- Frustum culling enabled\n\nDeliver a single self-contained HTML file with embedded CSS and JS using ES modules from CDN (unpkg.com/three@0.160.0/build/three.module.js). Include OrbitControls for interactive viewing.")

add_para(
    "Scoring was based on ten binary criteria: WebGLRenderer presence, Scene initialization, PerspectiveCamera, "
    "requestAnimationFrame loop, lighting (AmbientLight, DirectionalLight, or HemisphereLight), shadow mapping, "
    "3D geometry (BoxGeometry, SphereGeometry, or TorusGeometry), UI overlay (h1/h2/button), resize handler, "
    "and OrbitControls."
)

add_heading("2.3 Experiment 2: Ecommerce with Loop Engineering", level=2)
add_para(
    "Each model was tasked with building a complete ecommerce page for a fictional dog food brand called PawNutri. "
    "The prompt was delivered individually to each model. After the first generation, output was scored on 14 criteria. "
    "If missing features were found, a second iteration was run with specific feedback. The ecommerce prompt is "
    "reproduced below."
)

add_heading("Ecommerce Prompt (verbatim)", level=3)
add_code('''Return ONLY a complete HTML file. No explanation, no markdown formatting, no code blocks. Just raw HTML starting with <!DOCTYPE html>.

Create a complete ecommerce page for "PawNutri" dog food with:

1. Three.js 3D hero section: rotating dog food bag, floating kibble, warm professional lighting
2. Product grid fetching from Supabase: https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/products with headers apikey and Authorization using Bearer token
3. Shopping cart with localStorage add/remove/quantity/total
4. Checkout form: customer_name, customer_email, customer_address. Submit to https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/orders and https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/order_items
5. Confirmation message after successful order
6. Responsive warm design (browns, golds, creams)

CRITICAL: The product fetch AND order submission MUST use real API calls to Supabase. Do not hardcode products. Use this anon key: sb_publishable_HzadVL89uY9sfoBI4XnuHw_aieHtKG3
For Supabase headers use: apikey, Authorization Bearer, Content-Type application/json, Prefer: return=representation''')

add_para(
    "Scoring was based on 14 binary criteria: Three.js Scene, WebGLRenderer, PerspectiveCamera, requestAnimationFrame "
    "loop, lighting, product grid (map/forEach pattern), Supabase API fetch (actual fetch call to the products endpoint), "
    "cart UI (addToCart function), checkout form (fields: customer_name, email, address), order submission (fetch POST "
    "to orders/order_items), total calculation, localStorage persistence, responsive design (media query or flex-wrap), "
    "and order confirmation message."
)

# ===== 3. RESULTS =====
add_heading("3. Results", level=1)

add_heading("3.1 Experiment 1: Frontend Landing Page", level=2)

add_table(
    ["Model", "Vibe Score", "Technical Score", "Vibe Time (s)", "Tech Time (s)"],
    [
        ["Nemotron 3 Super", "9.0/10", "10.0/10", "29", "36"],
        ["Gemma 4", "9.0/10", "10.0/10", "16", "19"],
        ["MiniMax M3", "8.0/10", "0/10*", "36", "88"],
    ]
)
add_para("Table 2. Frontend landing page quality scores. *MiniMax M3 returned empty response for the technical prompt.", indent=False)

add_para(
    "Nemotron 3 Super and Gemma 4 both achieved 10/10 on the technical prompt, implementing all ten required features "
    "including OrbitControls. Their vibe-prompt scores of 9/10 were missing only OrbitControls. MiniMax M3 scored 8/10 "
    "on the vibe prompt, missing the resize handler and OrbitControls; its technical prompt response was empty — the "
    "model returned zero characters. Response times favored Gemma 4 (16-19 seconds) over Nemotron (29-36 seconds) "
    "and MiniMax (36-88 seconds)."
)

add_heading("3.2 Experiment 2: Ecommerce Store", level=2)

add_table(
    ["Model", "Iterations", "Best Score", "Missing Features", "API Calls Verified"],
    [
        ["Nemotron 3 Super", "2 (9.3 → 10.0)", "10.0/10", "None", "Yes — products, orders, order_items"],
        ["Gemma 4", "1", "10.0/10", "None", "Yes — products, orders, order_items"],
        ["MiniMax M3", "1", "10.0/10", "None", "Yes — products, orders, order_items"],
    ]
)
add_para("Table 3. Ecommerce loop engineering results.", indent=False)

add_para(
    "All three models achieved perfect 10/10 scores on the ecommerce task. Every model implemented real Supabase API "
    "calls for product fetching, order creation, and order_items insertion, using the provided Supabase URL and anon "
    "key with correct authentication headers. All models also included the Three.js 3D hero section, a product grid, "
    "shopping cart with localStorage, checkout form with three fields, total calculation, responsive design, and "
    "order confirmation messages."
)
add_para(
    "Nemotron 3 Super required two iterations: the first scored 9.3/10 (missing order submission to the database), "
    "and the second achieved 10.0/10 after specific feedback. Gemma 4 and MiniMax M3 both achieved 10.0/10 on the "
    "first attempt. Notably, the MiniMax M3 result contradicts an earlier batch evaluation where the same model "
    "failed to produce valid HTML. The difference is attributed to using the generate API instead of the chat API, "
    "which returned empty responses for MiniMax."
)

add_heading("3.3 Response Time Comparison", level=2)

add_table(
    ["Model", "Frontend Vibe", "Frontend Technical", "Ecommerce Iter 1"],
    [
        ["Nemotron 3 Super", "29s", "36s", "68s"],
        ["Gemma 4", "16s", "19s", "28s"],
        ["MiniMax M3", "36s", "88s (failed)", "59s"],
    ]
)
add_para("Table 4. Generation times in seconds across all tasks.", indent=False)

add_para(
    "Gemma 4 was the fastest model across all tasks, generating complex ecommerce pages in 28 seconds — roughly "
    "half the time of Nemotron 3 Super (68 seconds) and MiniMax M3 (59 seconds). Nemotron's ecommerce iteration 2 "
    "took 103 seconds, possibly due to the added context of the iteration 1 output being included in the prompt."
)

# ===== 4. DISCUSSION =====
add_heading("4. Discussion", level=1)

add_para(
    "The results of this study differ notably from earlier batch evaluations of the same models. In prior runs, "
    "all models plateaued at 8.6/10 on the ecommerce task and consistently failed to implement real Supabase API "
    "calls. In this one-by-one evaluation, all three models achieved perfect scores with verified API integration. "
    "This discrepancy highlights a critical methodology lesson: batch evaluation pipelines introduce failure modes "
    "that can dramatically underestimate model capability."
)
add_para(
    "Several specific factors contributed to pipeline errors in the batch approach. First, MiniMax M3 returned "
    "empty responses via the chat API but produced valid output via the generate API — a difference invisible "
    "in batch processing. Second, response parsing in batch mode used fragile regex patterns that could miss "
    "valid HTML, particularly when models used template literals or variable-based URL construction. Third, "
    "API timeouts in batch mode could truncate responses, especially for larger models like Nemotron 3 Super "
    "whose generation time exceeded 100 seconds for complex tasks."
)
add_para(
    "The prompt format itself was also critical. The batch ecommerce prompt was a dense, multi-paragraph "
    "specification that may have overwhelmed models. The one-by-one prompt was more structured, used imperative "
    "language (\"Return ONLY a complete HTML file\"), and explicitly warned against hardcoding data. These "
    "prompt engineering differences likely contributed to the improved results."
)
add_para(
    "All three models demonstrated the ability to generate working full-stack code with real API calls. This "
    "is a more encouraging finding than previous work suggested. However, the fragility of the results — "
    "sensitive to API endpoint choice, prompt formatting, and iteration strategy — means that practitioners "
    "cannot yet rely on these models to produce correct database-connected code without careful prompt "
    "engineering and output verification."
)

# ===== 5. CONCLUSION =====
add_heading("5. Conclusion", level=1)

add_para(
    "This study evaluated three cloud LLMs on frontend and full-stack ecommerce code generation, with models "
    "run individually and outputs verified manually. The key findings are as follows."
)
add_para(
    "First, all three models (Nemotron 3 Super, Gemma 4, MiniMax M3) achieved perfect 10/10 scores on the "
    "ecommerce task, generating working HTML pages with Three.js 3D scenes, product catalogs fetched from a "
    "live Supabase API, shopping carts with localStorage, checkout forms, and order submission to the database. "
    "This demonstrates that current cloud LLMs can generate real full-stack web applications when prompts are "
    "carefully constructed."
)
add_para(
    "Second, Gemma 4 was the fastest and most reliable model, generating perfect ecommerce code in 28 seconds "
    "on the first attempt. Nemotron 3 Super required a second iteration to fix a missing order submission feature. "
    "MiniMax M3 produced perfect output on the first attempt but only via the generate API endpoint."
)
add_para(
    "Third, batch evaluation methodologies can significantly underestimate model capability. Pipeline artifacts — "
    "wrong API endpoint selection, fragile HTML extraction, response parsing errors — can cause working code to "
    "be scored as failed. Future evaluations should consider running models individually with manual verification "
    "as a complement to automated batch scoring."
)
add_para(
    "Fourth, the exact format of the prompt matters enormously. The instruction \"Return ONLY a complete HTML "
    "file\" with explicit warnings against placeholder data significantly improved output quality. Models that "
    "produced placeholder data with verbose prompts generated real API calls when given concise, imperative "
    "instructions."
)
add_para(
    "For practitioners, the recommendation is clear: use Gemma 4 for the best speed-quality balance on web "
    "code generation tasks. Structure prompts as imperative commands with explicit constraints. Always verify "
    "that API calls use real endpoints rather than placeholder data. And if a model fails on a given API endpoint, "
    "try an alternative endpoint before concluding the model is incapable."
)

# ===== PROMPTS REFERENCE =====
add_heading("Appendix A: Complete Prompt Reference", level=1)
add_para(
    "All prompts used in this study are reproduced below for exact replication. Prompts are shown verbatim, "
    "including whitespace and formatting."
)

add_heading("A.1 Frontend Vibe Prompt", level=2)
add_code("Create a 3D dog food brand landing page using Three.js. Make it look professional and modern. It should have:\n\n1. A 3D scene with a rotating dog food bag or kibble pieces floating around\n2. Nice lighting and colors (warm tones, pet-friendly feel)\n3. A headline \"Premium Dog Food\" with a subtitle\n4. A \"Shop Now\" button\n5. The background should look nice with some atmospheric effects\n\nMake it a complete HTML file that works when opened in a browser. Keep it simple but impressive visually.")

add_heading("A.2 Frontend Technical Prompt", level=2)
add_code("Build a complete, production-quality 3D dog food landing page using Three.js (r152+ via CDN/ESM). Technical specifications:\n\n**Scene Setup:**\n- PerspectiveCamera (fov: 45, near: 0.1, far: 1000) positioned at (0, 2, 8)\n- WebGLRenderer with antialiasing, shadow mapping enabled (PCFSoftShadowMap), toneMapping: ACESFilmicToneMapping, toneMappingExposure: 1.2\n- Renderer size: window.innerWidth x window.innerHeight, pixelRatio: Math.min(window.devicePixelRatio, 2)\n\n**Lighting:**\n- AmbientLight: intensity 0.4\n- DirectionalLight: position (5, 10, 7), intensity 1.0, castShadow: true, shadow map size 2048x2048\n- HemisphereLight: sky #87CEEB, ground #3e7a3e, intensity 0.6\n\n**3D Objects:**\n- A dog food bag created with BoxGeometry(1.5, 2.2, 0.8), rounded via BufferGeometry custom vertices or use RoundedBoxGeometry from the examples\n- Texture: generate a canvas texture with dog food brand colors (#8B4513 brown, #DAA520 gold, #2E8B57 green) and text \"PREMIUM DOG FOOD\"\n- kibble pieces: 30-50 small TorusGeometry or SphereGeometry objects scattered around, slowly orbiting\n- A ground plane: CircleGeometry(4, 32) receiving shadows\n\n**Animation Loop:**\n- Dog food bag: smooth Y-axis rotation (0.3 rad/s), gentle bobbing (sine wave, amplitude 0.05, frequency 0.5)\n- Kibble pieces: orbit around center at varying radii (2-3.5), speeds (0.1-0.3 rad/s), with individual Y-axis oscillation\n- All animations via requestAnimationFrame, delta-time-based for consistent speed\n\n**UI Overlay:**\n- HTML/CSS overlay with flexbox centered layout\n- Header: \"PREMIUM DOG FOOD\" in bold serif font (#2C1810), text-shadow for depth, 2.5rem size\n- Subheader: \"Crafted with love. Backed by science.\" in lighter weight\n- CTA button: rounded, gradient background (#DAA520 to #8B4513), hover glow effect, box-shadow transition\n- Z-index management: overlay above canvas\n\n**Responsive:**\n- Handle window resize: update camera aspect ratio, renderer size\n- Mobile: reduce FOV to 50, reposition camera to (0, 3, 10)\n- Touch: device orientation or pointer lock optional\n\n**Performance:**\n- Use InstancedMesh for kibble if more than 50\n- Dispose geometries and materials on scene unmount\n- Frustum culling enabled\n\nDeliver a single self-contained HTML file with embedded CSS and JS using ES modules from CDN (unpkg.com/three@0.160.0/build/three.module.js). Include OrbitControls for interactive viewing.")

add_heading("A.3 Ecommerce Prompt (Initial)", level=2)
add_code('''Return ONLY a complete HTML file. No explanation, no markdown formatting, no code blocks. Just raw HTML starting with <!DOCTYPE html>.

Create a complete ecommerce page for "PawNutri" dog food with:

1. Three.js 3D hero section: rotating dog food bag, floating kibble, warm professional lighting
2. Product grid fetching from Supabase: https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/products with headers apikey and Authorization using Bearer token
3. Shopping cart with localStorage add/remove/quantity/total
4. Checkout form: customer_name, customer_email, customer_address. Submit to https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/orders and https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/order_items
5. Confirmation message after successful order
6. Responsive warm design (browns, golds, creams)

CRITICAL: The product fetch AND order submission MUST use real API calls to Supabase. Do not hardcode products. Use this anon key: sb_publishable_HzadVL89uY9sfoBI4XnuHw_aieHtKG3
For Supabase headers use: apikey, Authorization Bearer, Content-Type application/json, Prefer: return=representation''')

add_heading("A.4 Ecommerce Prompt (Iteration 2, Nemotron only)", level=2)
add_code("(Same as iteration 1, with the following added at the end:\n\n\"PREVIOUS VERSION WAS MISSING: order submission to Supabase. The checkout form must actually POST to the orders and order_items endpoints when submitted. Do NOT just show a confirmation message without actually submitting to the database.\")")

# ===== REFERENCES =====
add_heading("References", level=1)
refs = [
    "Chen, M., Tworek, J., Jun, H., Yuan, Q., Pinto, H. P. O., Kaplan, J., ... & Zaremba, W. (2021). Evaluating Large Language Models Trained on Code. arXiv:2107.03374.",
    "Google DeepMind. (2026). Gemma 4: Technical Report.",
    "MiniMax. (2026). MiniMax M3: Coding & Agentic Frontier.",
    "NVIDIA. (2026). Nemotron 3 Super: Open Hybrid Mamba-Transformer MoE. NVIDIA Research.",
    "Ollama. (2026). Cloud Models Documentation. https://docs.ollama.com/cloud.",
    "Team Qwen. (2024). Qwen2.5-Coder: Technical Report. arXiv:2409.12186.",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.hanging_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

out = Path.home() / "Desktop" / "LLM_Verified_Research_Paper.docx"
doc.save(str(out))
print(f"Paper saved: {out}")
