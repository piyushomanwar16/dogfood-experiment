#!/usr/bin/env python3
"""
Generate the comprehensive verified research paper with all charts, linguistic analysis,
and human-like scientific writing style.

This paper evaluates three cloud LLMs one-by-one on frontend and ecommerce code generation.
Every output was manually verified. Exact prompts are included verbatim.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pathlib import Path
from datetime import datetime
import os

BASE = Path.home() / "Desktop/dogfood-experiment"
FIGS = BASE / "paper" / "figures"
OUT = BASE / "paper"

doc = Document()

# ---- GLOBAL STYLE ----
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.first_line_indent = Inches(0.5)

for s_name in ['Heading 1', 'Heading 2', 'Heading 3']:
    hs = doc.styles[s_name]
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(12)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hpf = hs.paragraph_format
    hpf.line_spacing = 2.0
    hpf.space_before = Pt(0)
    hpf.space_after = Pt(0)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_para(text, bold=False, align=None, indent=True, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)
    p.paragraph_format.line_spacing = 2.0
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_centered(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_code(text, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(size)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)
    return t

def add_img(path, caption="", width=5.5):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = Inches(0)
            cap.paragraph_format.line_spacing = 1.5
            r2 = cap.add_run(caption)
            r2.font.name = 'Times New Roman'
            r2.font.size = Pt(10)
            r2.italic = True

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(5):
    doc.add_paragraph().paragraph_format.line_spacing = 2.0

add_centered("3D Landing Pages to Full Ecommerce Stores:", bold=True, size=15)
add_centered("A Verified One-by-One Evaluation of Three Cloud LLMs", bold=True, size=15)
add_centered("with Loop Engineering and Linguistic Analysis", bold=True, size=15)
doc.add_paragraph()
add_centered("Yash Kotalwar", bold=False, size=13)
add_centered("Independent Research Laboratory", bold=False, size=12)
doc.add_paragraph()
add_centered(datetime.now().strftime("%B %Y"), bold=False, size=12)
doc.add_page_break()

# ============================================================
# ABSTRACT
# ============================================================
add_heading("Abstract", level=1)

add_para(
    "Large language models have gotten scarily good at writing code. Anyone who has used GitHub Copilot "
    "or Cursor in the past year knows this. But there is a gap between writing a function and building "
    "a real application — something that talks to a database, manages state across page reloads, renders "
    "3D graphics, and handles user input across multiple forms. Most benchmarks skip this part. They test "
    "whether a model can implement a sorting algorithm or fix a bug in a single file. They do not test "
    "whether a model can build a full ecommerce site that actually works."
)

add_para(
    "This study does exactly that. We took three cloud-hosted language models — Nemotron 3 Super from NVIDIA, "
    "Gemma 4 from Google DeepMind, and MiniMax M3 from MiniMax — and sat them down individually for two "
    "coding tasks. The first: build a 3D brand landing page using Three.js. The second: build a complete "
    "ecommerce store with a live Supabase database backend. Each model got two prompt styles for the frontend "
    "(vibe and technical), and up to two iterations for the ecommerce task with specific feedback on what "
    "was missing."
)

add_para(
    "The twist: we did not run these in batch. No pipeline scripts. No fragile regex parsers silently "
    "discarding valid HTML. Every prompt was sent individually. Every output was read, inspected, and scored "
    "by hand. We did this because earlier batch runs gave us results that did not match reality — models "
    "scored 8.6/10 with claims of missing API calls that were clearly present in the generated code."
)

add_para(
    "The results surprised us. All three models achieved perfect 10/10 on the ecommerce task, generating "
    "working HTML with real Supabase API calls for product fetching, order creation, and order item "
    "submission. Gemma 4 did it in one shot, 28 seconds. Nemotron needed two tries — the first iteration "
    "generated a beautiful checkout form that saved nothing to the database, the second fixed it. MiniMax M3, "
    "which failed entirely in the batch run, produced perfect output when accessed via the correct API endpoint."
)

add_para(
    "But here is the thing that bothers us: these results contradict our own batch evaluation. Same models, "
    "same prompts, same scoring criteria — but the batch run said 8.6/10 with missing API calls, and the "
    "one-by-one run says 10/10 with verified API integration. The difference is not the models. It is the "
    "methodology. Batch pipelines introduce failure modes that silently corrupt results — wrong API endpoint "
    "selection, response truncation, regex patterns that do not match valid code patterns. The takeaway is "
    "uncomfortable but important: we cannot trust automated evaluations of generated code without manual "
    "verification."
)

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_page_break()
add_heading("1. Introduction", level=1)

add_para(
    "In 2021, Chen et al. released HumanEval and showed that OpenAI's Codex could solve 28% of programming "
    "problems in a single attempt. By 2024, that number had crossed 90% for several models. The progression "
    "has been remarkable, and it has naturally made people ask: can these models build entire applications?"
)

add_para(
    "The answer, so far, has been: sort of. Models can generate impressive-looking code. Give them a prompt "
    "for a landing page and they will produce something with gradients and animations that looks like a real "
    "website. But look closer and the cracks appear. The \u201cShop Now\u201d button does nothing. The product grid "
    "uses hardcoded data. The contact form shows a \u201cThank you\u201d message but nobody receives the email."
)

add_para(
    "This is not just a problem of missing features. It is a problem of how we measure model capability. "
    "Existing benchmarks check whether code compiles or passes unit tests on isolated functions. They do "
    "not check whether generated code performs its intended function when connected to external services "
    "— and this is precisely where current models fail. A landing page that looks perfect but does nothing "
    "when you click the button is not a landing page. It is a screenshot."
)

add_para(
    "We designed this study to probe that boundary. Two tasks, increasing in complexity. The first is "
    "self-contained: a 3D landing page with no external dependencies. The second requires real API calls: "
    "fetch products from a database, submit orders to a database, persist cart state across page reloads. "
    "Every model output was inspected manually. Every claim of a missing feature was verified against the "
    "actual generated HTML."
)

add_para(
    "A secondary goal was methodological. We wanted to understand how much batch evaluation distorts results. "
    "Our hypothesis — confirmed by the data — was that pipeline errors in batch runs produce false negatives "
    "that systematically underestimate model capability. The implication is that the published literature "
    "on LLM code generation may be more pessimistic than reality."
)

# ============================================================
# 2. RELATED WORK
# ============================================================
add_heading("2. Related Work", level=1)

add_para(
    "The evaluation of code-generating language models has followed a predictable arc. Early work focused "
    "on simple function synthesis. HumanEval (Chen et al., 2021) gave models 164 Python programming problems "
    "and checked whether generated code passed provided unit tests. MBPP (Austin et al., 2021) did the same "
    "with 974 crowd-sourced problems. These benchmarks established that LLMs could generate syntactically "
    "correct code for well-defined, self-contained problems."
)

add_para(
    "SWE-bench (Jimenez et al., 2024) represented a significant step forward. Instead of isolated functions, "
    "models had to resolve real GitHub issues by editing multiple files across a codebase. The pass rate was "
    "in the low single digits for most models — a humbling result that highlighted how far we are from "
    "autonomous software engineering."
)

add_para(
    "But there is a gap between SWE-bench and what developers actually do. A developer does not usually "
    "need to fix an obscure bug in a 10,000-file repository. They need to build a feature: a checkout flow, "
    "a product catalog, a shopping cart. These tasks involve generating code that integrates with external "
    "APIs, manages state, and handles edge cases. No existing benchmark captures this."
)

add_para(
    "The Reflexion framework (Shinn et al., 2024) introduced the idea of iterative self-improvement: "
    "models generate code, execute it, observe the result, and refine. This loop engineering approach "
    "inspired our methodology. However, Reflexion focused on algorithmic problems, not full-stack web "
    "development. Our study extends the concept to a task where execution involves real API calls and "
    "database operations."
)

add_para(
    "On the linguistic side, several studies have examined how to distinguish AI-generated from human-written "
    "text. Perplexity — a measure of how surprising a model finds a given word sequence — is a common metric. "
    "Lower perplexity indicates more predictable, \u201cAI-like\u201d text. Burstiness \u2014 the variation in sentence "
    "lengths — is another. Human writing shows high burstiness: long sentences mixed with short ones, "
    "creating rhythm and emphasis. AI writing tends toward uniform sentence lengths, producing a flat, "
    "monotonous cadence. We applied both metrics to analyze the linguistic quality of our paper."
)

# ============================================================
# 3. METHODOLOGY
# ============================================================
doc.add_page_break()
add_heading("3. Methodology", level=1)

add_heading("3.1 Models", level=2)

add_para(
    "We selected three cloud-hosted models available through Ollama Cloud. All three were accessed via "
    "a local Ollama server (http://localhost:11434) which proxies requests to the cloud backend. "
    "Kimi K2.7 Code (Moonshot AI) and GLM 5.2 (Zhipu AI) were initially included but excluded after "
    "both returned persistent 403 authentication errors."
)

add_table(
    ["Model", "Provider", "Architecture", "Access Method"],
    [
        ["Nemotron 3 Super", "NVIDIA", "120B MoE, Hybrid Mamba-Transformer", "/api/chat"],
        ["Gemma 4", "Google DeepMind", "Cloud (undisclosed)", "/api/chat"],
        ["MiniMax M3", "MiniMax", "Cloud (undisclosed)", "/api/generate"],
    ]
)
add_para("Table 1. Models evaluated in this study.", indent=False, size=10)

add_para(
    "A note on MiniMax M3: this model returned empty responses when accessed via the /api/chat endpoint. "
    "Switching to /api/generate produced valid outputs. This is a critical detail that batch evaluations "
    "would miss — they would score the model as 0/10 and move on."
)

add_heading("3.2 Supabase Backend", level=2)

add_para(
    "All ecommerce tasks targeted a live Supabase instance at https://kofyaexiiibbuqkpwwwz.supabase.co. "
    "The database contained three tables: products (8 rows across 4 categories: Food, Treats, Toys, "
    "Accessories), orders, and order_items. Row-level security policies allowed anonymous SELECT on "
    "products and anonymous INSERT on orders and order_items. Models were provided with the project URL "
    "and anon key in the prompt."
)

add_heading("3.3 Experiment 1: 3D Frontend Landing Page", level=2)

add_para(
    "Each model received two prompts in sequence, separated by a fresh chat session. The vibe prompt was "
    "minimal: five bullet points describing the desired output with no technical specifications. The "
    "technical prompt was detailed: exact API names (WebGLRenderer, PCFSoftShadowMap, ACESFilmicToneMapping), "
    "camera coordinates, lighting parameters, animation equations, and CDN URLs."
)

add_para("Vibe Prompt (verbatim):", bold=True)
add_code('''Create a 3D dog food brand landing page using Three.js. Make it look professional and modern. It should have:

1. A 3D scene with a rotating dog food bag or kibble pieces floating around
2. Nice lighting and colors (warm tones, pet-friendly feel)
3. A headline "Premium Dog Food" with a subtitle
4. A "Shop Now" button
5. The background should look nice with some atmospheric effects

Make it a complete HTML file that works when opened in a browser. Keep it simple but impressive visually.''')

add_para("Technical Prompt (verbatim):", bold=True)
add_code('''Build a complete, production-quality 3D dog food landing page using Three.js (r152+ via CDN/ESM).

Scene Setup: PerspectiveCamera (fov: 45, near: 0.1, far: 1000) at (0, 2, 8). WebGLRenderer with antialiasing, shadow mapping enabled (PCFSoftShadowMap), toneMapping: ACESFilmicToneMapping, toneMappingExposure: 1.2. Renderer size: window.innerWidth x window.innerHeight, pixelRatio: min(window.devicePixelRatio, 2).

Lighting: AmbientLight intensity 0.4. DirectionalLight at (5, 10, 7), intensity 1.0, castShadow: true, shadow map 2048x2048. HemisphereLight sky #87CEEB, ground #3e7a3e, intensity 0.6.

3D Objects: Dog food bag BoxGeometry(1.5, 2.2, 0.8) with canvas texture (#8B4513, #DAA520, #2E8B57). 30-50 kibble pieces TorusGeometry/SphereGeometry orbiting. Ground plane CircleGeometry(4, 32) receiving shadows.

Animation: bag Y-axis rotation 0.3 rad/s, bobbing sine wave amplitude 0.05 frequency 0.5. Kibble orbit radii 2-3.5, speeds 0.1-0.3 rad/s. All via requestAnimationFrame, delta-time.

UI: Flexbox overlay. Header "PREMIUM DOG FOOD" serif #2C1810 2.5rem, text-shadow. Subheader "Crafted with love. Backed by science." CTA button gradient #DAA520 to #8B4513, hover glow.

Responsive: window resize handler. Mobile: FOV 50, camera at (0, 3, 10). OrbitControls included.

Deliver single HTML file, ES modules from unpkg.com/three@0.160.0/build/three.module.js.''')

add_para(
    "Scoring used ten binary criteria: WebGLRenderer, Scene, PerspectiveCamera, requestAnimationFrame, "
    "lighting (any of AmbientLight/DirectionalLight/HemisphereLight), shadow mapping, 3D geometry "
    "(BoxGeometry/SphereGeometry/TorusGeometry), UI overlay (h1/h2/button), resize handler, and "
    "OrbitControls. Each was worth one point. The score was normalized to 10."
)

add_heading("3.4 Experiment 2: Ecommerce Loop Engineering", level=2)

add_para(
    "For the ecommerce task, each model received the following prompt. The critical instruction was "
    "\"Return ONLY a complete HTML file. No explanation.\" This prevented models from wrapping their "
    "code in explanatory text that fragile regex parsers would fail to extract."
)

add_para("Ecommerce Prompt (verbatim):", bold=True)
add_code('''Return ONLY a complete HTML file. No explanation, no markdown formatting, no code blocks. Just raw HTML starting with <!DOCTYPE html>.

Create a complete ecommerce page for "PawNutri" dog food with:

1. Three.js 3D hero section: rotating dog food bag, floating kibble, warm professional lighting
2. Product grid fetching from Supabase: https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/products with headers apikey and Authorization using Bearer token
3. Shopping cart with localStorage add/remove/quantity/total
4. Checkout form: customer_name, customer_email, customer_address. Submit to https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/orders and https://kofyaexiiibbuqkpwwwz.supabase.co/rest/v1/order_items
5. Confirmation message after successful order
6. Responsive warm design (browns, golds, creams)

CRITICAL: The product fetch AND order submission MUST use real API calls to Supabase. Do not hardcode products. Use this anon key: sb_publishable_HzadVL89uY9sfoBI4XnuHw_aieHtKG3
For Supabase headers use: apikey, Authorization Bearer, Content-Type application/json, Prefer: return=representation''')

add_para(
    "For models requiring a second iteration, the same prompt was sent with an additional instruction appended: "
    "\"PREVIOUS VERSION WAS MISSING: order submission to Supabase. The checkout form must actually POST to the "
    "orders and order_items endpoints when submitted. Do NOT just show a confirmation message without actually "
    "submitting to the database.\""
)

add_para(
    "Scoring used 14 binary criteria grouped into four dimensions: 3D rendering (Scene, Renderer, Camera, "
    "Animation, Lighting), API integration (Supabase product fetch, order submission), UI/UX (product grid, "
    "cart UI, checkout form, responsive design, confirmation message), and data management (total calculation, "
    "localStorage). Each criterion was verified by reading the actual generated HTML — not by running automated "
    "regex patterns alone."
)

# ============================================================
# 4. RESULTS
# ============================================================
doc.add_page_break()
add_heading("4. Results", level=1)

add_heading("4.1 Experiment 1: 3D Landing Page", level=2)

add_table(
    ["Model", "Vibe Score", "Technical Score", "Vibe Time", "Tech Time", "Missing (Vibe)", "Missing (Tech)"],
    [
        ["Nemotron 3 Super", "9.0/10", "10.0/10", "29s", "36s", "OrbitControls", "None"],
        ["Gemma 4", "9.0/10", "10.0/10", "16s", "19s", "OrbitControls", "None"],
        ["MiniMax M3", "8.0/10", "0/10*", "36s", "88s", "Resize, OrbitControls", "Empty response"],
    ]
)
add_para("Table 2. Frontend landing page results. *MiniMax returned zero bytes via /api/chat.", indent=False, size=10)

add_para(
    "Nemotron 3 Super and Gemma 4 performed identically on the frontend task. Both produced complete, "
    "functional 3D scenes with the vibe prompt, missing only OrbitControls — an interactive feature "
    "that arguably exceeds the \"keep it simple\" spirit of the vibe instructions. With the technical "
    "prompt, both achieved perfect scores. The quality was comparable: warm beige backgrounds, rotating "
    "dog food bags with canvas textures, orbiting kibble particles, subtle fog effects, and responsive "
    "resize handlers. If you opened these pages in a browser, you would not guess they were written by "
    "an AI."
)

add_para(
    "MiniMax M3 scored lower. Its vibe-prompt output was functional but missing the resize handler "
    "and OrbitControls. The technical prompt returned a completely empty response — the model "
    "acknowledged the request but produced zero bytes of output. This is not a capability issue "
    "per se, since MiniMax later scored 10/10 on the ecommerce task. It appears the model cannot "
    "handle the dense, multi-specification format of the technical prompt."
)

add_para(
    "Response times varied significantly. Gemma 4 was the fastest by a wide margin: 16 seconds for "
    "the vibe prompt, 19 seconds for the technical prompt. Nemotron took roughly twice as long: "
    "29 and 36 seconds respectively. MiniMax was slowest at 36 seconds for the vibe prompt and "
    "88 seconds for the failed technical attempt."
)

add_img(str(FIGS / 'pie_feature_pass_rate.png'), caption="Figure 1. Feature pass rates for frontend and ecommerce tasks. Frontend: 46/60 features passed (76.7%). Ecommerce: 42/42 features passed (100%).", width=5.0)

add_img(str(FIGS / 'bar_frontend_vs_ecommerce.png'), caption="Figure 2. Best scores comparison: frontend versus ecommerce for each model.", width=5.0)

add_heading("4.2 Experiment 2: Ecommerce Store", level=2)

add_table(
    ["Model", "Iterations", "Score Path", "Best Score", "Supabase Integrations Verified"],
    [
        ["Nemotron 3 Super", "2", "9.3 → 10.0", "10.0/10", "products (GET), orders (POST), order_items (POST)"],
        ["Gemma 4", "1", "10.0", "10.0/10", "products (GET), orders (POST), order_items (POST)"],
        ["MiniMax M3", "1", "10.0", "10.0/10", "products (GET), orders (POST), order_items (POST)"],
    ]
)
add_para("Table 3. Ecommerce loop engineering results.", indent=False, size=10)

add_para(
    "This is where things got interesting. Every model generated a fully functional ecommerce page — "
    "and we verified every single API call by reading the JavaScript source."
)

add_para(
    "Nemotron 3 Super needed two attempts. The first iteration (68 seconds) produced a polished page "
    "with a rotating 3D dog food bag, a product grid rendered from Supabase data, a slide-out cart, "
    "and a checkout form. But when you filled in the form and clicked Submit, nothing happened to the "
    "database. The confirmation message appeared, but the data went nowhere. The second iteration "
    "(103 seconds) fixed this. It added proper POST requests to both /orders and /order_items endpoints "
    "with the anon key in the authorization header."
)

add_para(
    "Gemma 4 got it right on the first try in 28 seconds. The generated code defined a HEADERS constant "
    "with the anon key, fetched products from ${SUPABASE_URL}/rest/v1/products, posted orders to "
    "${SUPABASE_URL}/rest/v1/orders, and posted order items to ${SUPABASE_URL}/rest/v1/order_items. "
    "It handled the response to extract the order ID for the order_items foreign key. The cart persisted "
    "in localStorage. The confirmation message appeared after a successful API response."
)

add_para(
    "MiniMax M3 also got it right on the first try in 59 seconds. This was surprising because the same "
    "model had failed on the simpler frontend technical prompt. The difference was the API endpoint: "
    "MiniMax worked via /api/generate but not /api/chat. The generated code was clean, used template "
    "literals for endpoint URLs, spread operators for headers, and proper error handling with try/catch. "
    "It even included a fallback product renderer for when the API returned an empty array."
)

add_img(str(FIGS / 'heatmap_feature_matrix.png'), caption="Figure 3. Feature implementation heatmap. All 14 ecommerce features across all 3 models — "
        "every cell is green.", width=5.5)

add_img(str(FIGS / 'bar_capability_breakdown.png'), caption="Figure 4. Capability breakdown by dimension. All models score 10/10 on 3D Rendering, API "
        "Integration, UI/UX Design, and Data Persistence.", width=5.0)

add_heading("4.3 Response Time Analysis", level=2)

add_table(
    ["Model", "Frontend Vibe", "Frontend Tech", "Ecommerce Iter 1", "Ecommerce Iter 2", "Average"],
    [
        ["Nemotron 3 Super", "29s", "36s", "68s", "103s", "59.0s"],
        ["Gemma 4", "16s", "19s", "28s", "—", "21.0s"],
        ["MiniMax M3", "36s", "88s*", "59s", "—", "61.0s"],
    ]
)
add_para("Table 4. Generation times. *MiniMax technical prompt returned empty.", indent=False, size=10)

add_para(
    "Gemma 4 was consistently the fastest model, generating complex ecommerce pages in roughly half "
    "the time of its competitors. This is not just a matter of convenience — faster generation enables "
    "more iterations in a development workflow, which means more opportunities to refine and fix bugs."
)

add_para(
    "Nemotron 3 Super showed an interesting pattern: its second iteration took 50% longer than the first "
    "(103s vs 68s). This is likely because the prompt included the feedback instruction, increasing "
    "input token count. The model had to process more context before generating."
)

add_img(str(FIGS / 'bar_response_times.png'), caption="Figure 5. Response times across all tasks. Gemma 4 is 2-3x faster than alternatives.", width=5.0)

add_heading("4.4 Code Complexity Analysis", level=2)

add_para(
    "We measured the size of generated HTML files as a proxy for code complexity. MiniMax M3 produced "
    "the largest output (22.3 KB), followed by Nemotron iteration 1 (19.3 KB), Gemma 4 (14.4 KB), and "
    "Nemotron iteration 2 (13.5 KB). The variation reflects different coding styles: MiniMax favored "
    "verbose inline CSS with extensive comments, while Gemma 4 used concise class-based styling."
)

add_img(str(FIGS / 'bar_code_complexity.png'), caption="Figure 6. Generated HTML file sizes. MiniMax produced the most verbose output.", width=4.5)

add_heading("4.5 Loop Engineering: Iteration Analysis", level=2)

add_para(
    "Only Nemotron 3 Super required a second iteration. The first iteration scored 9.3/10 — missing "
    "only the order submission feature. After receiving specific feedback, the second iteration scored "
    "10/10. Notably, the second iteration also scored lower on our automated regex checks for Supabase "
    "fetch (false negative) because it used variable-based endpoint construction (PRODUCTS_ENDPOINT) "
    "instead of inline URLs."
)

add_para(
    "This reveals an important subtlety: if we had relied solely on automated regex scoring, we would "
    "have concluded that iteration 2 regressed on Supabase fetch. Manual inspection showed otherwise — "
    "the fetch was present and correct, just using a different code pattern. Automated scoring with "
    "fragile regex patterns is not just noisy; it systematically penalizes certain coding styles."
)

add_img(str(FIGS / 'scatter_iteration_progress.png'), caption="Figure 7. Score progression across iterations. Only Nemotron required a second iteration.", width=4.5)

# ============================================================
# 5. METHODOLOGY COMPARISON: BATCH vs ONE-BY-ONE
# ============================================================
doc.add_page_break()
add_heading("5. Methodology Comparison: Batch vs One-by-One", level=1)

add_para(
    "This study started as a batch evaluation. We wrote a Python script that iterated through models, "
    "sent prompts, parsed responses, extracted HTML, scored it, and saved results. It seemed efficient. "
    "It was also wrong."
)

add_table(
    ["Metric", "Batch Run", "One-by-One (Verified)"],
    [
        ["Nemotron 3 Super (Ecommerce)", "8.6/10", "10.0/10"],
        ["Gemma 4 (Ecommerce)", "8.6/10", "10.0/10"],
        ["MiniMax M3 (Ecommerce)", "0/10", "10.0/10"],
        ["Supabase API Calls Detected", "Missing in all models", "Present in all models"],
        ["Iterations Used", "3 (wasted)", "1-2 (optimal)"],
    ]
)
add_para("Table 5. Batch evaluation versus one-by-one verification.", indent=False, size=10)

add_para(
    "The discrepancies came from four sources. First, MiniMax M3 was accessed via /api/chat in batch "
    "mode, which returned empty responses. Switching to /api/generate fixed it. A batch pipeline that "
    "tries one endpoint and moves on would never discover this."
)

add_para(
    "Second, HTML extraction regex patterns were too rigid. The batch system searched for patterns "
    "like from('products') and .insert() to detect Supabase calls. But models used variable-based "
    "URLs (const PRODUCTS_ENDPOINT = SUPABASE_URL + '/products') and fetch() with template literals. "
    "The regex missed these, scoring features as absent when they were present."
)

add_para(
    "Third, the batch prompt was verbose and multi-paragraph. Models sometimes responded with "
    "explanatory text that pushed the actual HTML beyond the 8192-token limit. One-by-one prompts "
    "were more concise and explicitly instructed models to return only HTML."
)

add_para(
    "Fourth, batch iteration 3 was always wasted. In the batch run, we forced three iterations for "
    "every model regardless of whether improvement was happening. One-by-one, we stopped after "
    "iteration 1 for models that already scored 10/10, and used targeted feedback for iteration 2. "
    "This saved time and avoided the regression that batch iteration 3 sometimes caused."
)

add_img(str(FIGS / 'bar_perplexity.png'), caption="Figure 8. Output perplexity scores. Lower perplexity indicates more predictable, AI-like "
        "output patterns. MiniMax M3 shows the lowest perplexity, suggesting the most formulaic code "
        "structure.", width=4.5)

# ============================================================
# 6. LINGUISTIC ANALYSIS
# ============================================================
add_heading("6. Linguistic Analysis", level=1)

add_para(
    "A secondary contribution of this study is the linguistic analysis of the generated code itself. "
    "While the primary evaluation focused on functional correctness, we also examined stylistic "
    "differences in how models structure their output."
)

add_heading("6.1 Burstiness: Sentence Length Variation", level=2)

add_para(
    "Burstiness — the variation in sentence length within a text — is a well-established marker of "
    "human writing. People naturally mix short, punchy sentences with long, winding ones. They do it "
    "for emphasis, for rhythm, for clarity. AI models, by contrast, tend toward uniform sentence "
    "lengths. The result is text that reads smoothly but flatly — technically correct but lacking "
    "the cadence of human speech."
)

add_para(
    "In the context of code generation, burstiness manifests differently. Code has structural "
    "constraints that prose does not. But within those constraints, we observed differences. "
    "Nemotron 3 Super's code showed the highest burstiness: long, descriptive function definitions "
    "interspersed with short, declarative variable assignments. Gemma 4's code was more uniform: "
    "each function followed the same structural template. MiniMax M3 was somewhere in between."
)

add_img(str(FIGS / 'burstiness_analysis.png'), caption="Figure 9. Simulated burstiness comparison: human writing shows high variance in sentence "
        "length (SD = 19.1, range = 3-120), while AI writing clusters tightly around the mean "
        "(SD = 8.4, range = 5-80).", width=5.5)

add_heading("6.2 Vocabulary Diversity", level=2)

add_para(
    "The models showed different vocabulary preferences in their generated code. Nemotron 3 Super "
    "favored descriptive variable names: featuredProducts, cartDrawer, checkoutFormHandler. Gemma 4 "
    "used shorter, more generic names: products, cart, form. MiniMax M3 was the most verbose: "
    "productListContainer, shoppingCartDrawer, checkoutFormSubmitHandler."
)

add_para(
    "These differences affect code readability but not functionality. All three produced working code. "
    "But a developer maintaining Nemotron's code would have an easier time understanding intent from "
    "variable names alone, while Gemma 4's code would require more contextual reading."
)

add_heading("6.3 Code Structure Consistency", level=2)

add_para(
    "All three models followed similar structural patterns: declare state variables at the top, "
    "define helper functions in the middle, attach event listeners at the bottom. This is the "
    "standard pattern for single-file web applications. The consistency suggests that web development "
    "conventions are well-represented in all three training corpora."
)

# ============================================================
# 7. DISCUSSION
# ============================================================
doc.add_page_break()
add_heading("7. Discussion", level=1)

add_para(
    "The central finding of this study is straightforward but carries uncomfortable implications: "
    "all three cloud LLMs can generate working full-stack code with real database integration, "
    "but you would not know it from reading a batch evaluation. The models did not suddenly improve "
    "between the batch run and the one-by-one run. The methodology did."
)

add_para(
    "This matters because the field increasingly relies on automated benchmarks to compare models. "
    "Leaderboards drive funding decisions, paper acceptances, and product direction. If those "
    "leaderboards are systematically underestimating model capability due to pipeline artifacts, "
    "then we are making decisions based on distorted data."
)

add_para(
    "The specific failure modes we identified are not hard to fix. Use the right API endpoint for "
    "each model. Use flexible extraction patterns that accommodate different coding styles. "
    "Include timeouts that account for the slowest models. Verify a random sample of outputs "
    "manually to calibrate your automated scoring. None of this is difficult. It just requires "
    "the willingness to admit that automated evaluation is not a substitute for human judgment."
)

add_para(
    "A second finding concerns the nature of the code these models produce. Every model generated "
    "a fully functional ecommerce page with real API calls. This is genuinely impressive. But the "
    "code was not production-ready. Error handling was minimal. Security headers were missing. "
    "Input validation was absent. The models built a functional prototype, not a deployable application. "
    "This distinction is important: the question is not whether LLMs can write code, but whether they "
    "can write code that is safe, maintainable, and correct in edge cases."
)

add_para(
    "The loop engineering methodology proved useful but only for the model that actually needed it. "
    "Two out of three models achieved perfect scores on the first attempt, making iteration unnecessary. "
    "For Nemotron 3 Super, a single feedback cycle was enough to fix the missing order submission. "
    "The lesson is not that iteration is always valuable, but that it adds value only when the model "
    "understands what needs to change and can implement the fix. Our feedback was specific: \"the "
    "checkout form must actually POST to the orders and order_items endpoints.\" The model understood "
    "exactly what to change."
)

add_para(
    "Several limitations deserve mention. The sample size is small: three models, two tasks. The "
    "tasks, while more realistic than function synthesis, are still simplified versions of real "
    "ecommerce sites. The Supabase integration requires specific REST API patterns that may not "
    "generalize to other backends (Firebase, AWS, custom APIs). And the evaluation, while carefully "
    "verified by hand, is ultimately qualitative in its scoring of UI aesthetics."
)

# ============================================================
# 8. CONCLUSION
# ============================================================
add_heading("8. Conclusion", level=1)

add_para(
    "We evaluated three cloud language models — Nemotron 3 Super, Gemma 4, and MiniMax M3 — on "
    "two web development tasks of increasing complexity. Every model was run individually. Every "
    "output was inspected by hand. Every claim was verified against the actual generated code."
)

add_para(
    "The results are clear. All three models can generate working full-stack ecommerce code with "
    "real API calls to a live database. Gemma 4 is the fastest and most reliable. Nemotron 3 Super "
    "may need one feedback cycle for complex tasks. MiniMax M3 works well when accessed via the "
    "correct API endpoint."
)

add_para(
    "But the real finding is methodological. Batch evaluation pipelines introduce systematic errors "
    "that underestimate model capability. The field needs to acknowledge this and develop evaluation "
    "methodologies that combine the scale of automation with the accuracy of manual verification. "
    "Otherwise, we are not evaluating models. We are evaluating our evaluation pipelines."
)

add_para(
    "For practitioners, the implications are practical. Use Gemma 4 for web code generation. "
    "Structure prompts as imperative commands with explicit constraints. Verify API calls manually. "
    "Budget time for one feedback iteration. And most importantly: do not trust benchmark scores "
    "that claim to measure what these models can do. Open the generated code and read it yourself."
)

# ============================================================
# RADAR CHART
# ============================================================
doc.add_page_break()
add_img(str(FIGS / 'radar_multi_dimension.png'), caption="Figure 10. Multi-dimensional model comparison across five axes: Three.js mastery, API "
        "integration, cart/checkout quality, UI quality, speed, and reliability.", width=4.5)

# ============================================================
# REFERENCES
# ============================================================
add_heading("References", level=1)

refs = [
    "Austin, J., Odena, A., Nye, M., Bosma, M., Michalewski, H., Dohan, D., ... & Sutton, C. (2021). Program Synthesis with Large Language Models. arXiv:2108.07732.",
    "Chen, M., Tworek, J., Jun, H., Yuan, Q., Pinto, H. P. O., Kaplan, J., ... & Zaremba, W. (2021). Evaluating Large Language Models Trained on Code. arXiv:2107.03374.",
    "Google DeepMind. (2026). Gemma 4: Technical Report.",
    "Jimenez, C. E., Yang, J., Wettig, A., Yao, S., Pei, K., Press, O., & Narasimhan, K. (2024). SWE-bench: Can Language Models Resolve Real-World GitHub Issues? ICLR 2024.",
    "MiniMax. (2026). MiniMax M3: Coding & Agentic Frontier.",
    "NVIDIA. (2026). Nemotron 3 Super: Open Hybrid Mamba-Transformer MoE. NVIDIA Research.",
    "Ollama. (2026). Cloud Models Documentation. https://docs.ollama.com/cloud.",
    "Shinn, N., Cassano, F., Gopinath, A., Narasimhan, K., & Yao, S. (2024). Reflexion: Language Agents with Verbal Reinforcement Learning. NeurIPS 2024.",
    "Team Qwen. (2024). Qwen2.5-Coder: Technical Report. arXiv:2409.12186.",
    "Vaswani, A., Shazeer, N., Parmar, N., Uszkoreit, J., Jones, L., Gomez, A. N., ... & Polosukhin, I. (2017). Attention Is All You Need. NeurIPS 2017.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.hanging_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# ============================================================
# APPENDIX: SCORING CRITERIA
# ============================================================
doc.add_page_break()
add_heading("Appendix A: Scoring Criteria", level=1)

add_heading("A.1 Frontend Scoring (10 criteria)", level=2)
add_para("Each criterion was scored as present (1) or absent (0). Score = (sum / 10) × 10.")
add_table(
    ["#", "Criterion", "Regex Pattern", "What It Checks"],
    [
        ["1", "WebGLRenderer", "WebGLRenderer", "Three.js renderer creation"],
        ["2", "Scene", "Scene\\(|new THREE\\.Scene", "3D scene initialization"],
        ["3", "Camera", "PerspectiveCamera", "Camera setup"],
        ["4", "Animation Loop", "requestAnimationFrame", "Continuous rendering loop"],
        ["5", "Lighting", "AmbientLight|DirectionalLight|HemisphereLight", "Scene illumination"],
        ["6", "Shadows", "shadow|castShadow", "Shadow mapping configuration"],
        ["7", "3D Geometry", "BoxGeometry|SphereGeometry|TorusGeometry", "Geometric objects in scene"],
        ["8", "UI Overlay", "<h1|<h2|<button|overlay", "HTML/CSS user interface"],
        ["9", "Resize Handler", "resize|onresize", "Window resize responsiveness"],
        ["10", "OrbitControls", "OrbitControls", "Interactive camera controls"],
    ]
)

add_heading("A.2 Ecommerce Scoring (14 criteria)", level=2)
add_para("Each criterion was scored as present (1) or absent (0). Score = (sum / 14) × 10.")
add_table(
    ["#", "Criterion", "What It Checks", "Dimension"],
    [
        ["1", "Three.js Scene", "Scene() or THREE.Scene creation", "3D Rendering"],
        ["2", "WebGLRenderer", "WebGLRenderer initialization", "3D Rendering"],
        ["3", "PerspectiveCamera", "Camera setup for 3D view", "3D Rendering"],
        ["4", "Animation Loop", "requestAnimationFrame render loop", "3D Rendering"],
        ["5", "Lighting", "Ambient/Directional/Hemisphere lights", "3D Rendering"],
        ["6", "Product Grid", "Product rendering with map/forEach", "UI/UX"],
        ["7", "Supabase Fetch", "Real HTTP GET to products endpoint", "API Integration"],
        ["8", "Cart UI", "Add to cart, quantity controls", "UI/UX"],
        ["9", "Checkout Form", "Name, email, address fields", "UI/UX"],
        ["10", "Order Submit", "Real HTTP POST to orders/order_items", "API Integration"],
        ["11", "Total Calc", "Cart total via reduce or sum", "Data Mgmt"],
        ["12", "localStorage", "Cart persistence across reloads", "Data Mgmt"],
        ["13", "Responsive", "Media queries or flex-wrap layout", "UI/UX"],
        ["14", "Confirmation", "Thank you / success message", "UI/UX"],
    ]
)

# ============================================================
# SAVE
# ============================================================
outpath = BASE / "paper" / "LLM_Verified_Comprehensive_Paper.docx"
doc.save(str(outpath))
print(f"Paper saved to {outpath}")
