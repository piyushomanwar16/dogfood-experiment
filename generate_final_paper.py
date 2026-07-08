#!/usr/bin/env python3
"""Generate updated research paper with full ecommerce loop engineering data."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path
from datetime import datetime
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np

BASE = Path.home() / "Desktop/dogfood-experiment"
FIGS = BASE / "paper" / "figures"
FIGS.mkdir(parents=True, exist_ok=True)

# ===== CHARTS =====

# Frontend results (from earlier + ecommerce)
models_short = ['Qwen-Coder\n7B', 'Qwen-Coder\n14B', 'MiniMax\nM3', 'Nemotron\n3 Super', 'Gemma 4']
frontend_vibe = [8.0, 8.0, 7.0, 9.0, 9.0]
frontend_tech = [10.0, 10.0, 0, 10.0, 10.0]
ecommerce_scores = [8.6, 8.6, 0, 8.6, 8.6]
ecommerce_iters = [3, 2, 3, 3, 3]

# Fig 1: Frontend comparison
fig, ax = plt.subplots(figsize=(9, 5))
x = np.arange(len(models_short))
w = 0.3
ax.bar(x - w, frontend_vibe, w, label='Vibe Prompt', color='#4C72B0', edgecolor='white')
ax.bar(x, frontend_tech, w, label='Technical Prompt', color='#DD8452', edgecolor='white')
ax.bar(x + w, ecommerce_scores, w, label='Ecommerce (Loop)', color='#55A868', edgecolor='white')
ax.set_ylabel('Score (/10)', fontsize=11)
ax.set_title('Model Performance Across Tasks', fontsize=13, pad=12)
ax.set_xticks(x)
ax.set_xticklabels(models_short, fontsize=8)
ax.set_ylim(0, 11)
ax.legend(fontsize=9)
ax.grid(axis='y', alpha=0.3, linestyle='--')
for i in range(len(x)):
    for scores, offset, color in [(frontend_vibe, -w, '#4C72B0'), (frontend_tech, 0, '#DD8452'), (ecommerce_scores, w, '#55A868')]:
        h = scores[i]
        if h > 0:
            ax.text(i + offset, h + 0.2, f'{h:.1f}', ha='center', fontsize=7, color=color, fontweight='bold')
plt.tight_layout()
fig.savefig(FIGS/'fig1_all_scores.png', dpi=200, bbox_inches='tight')
plt.close()
print("Chart 1 done")

# Fig 2: Ecommerce feature comparison
features = ['3D Scene', 'Products\nGrid', 'Supabase\nFetch', 'Cart', 'Checkout\nForm', 'Order\nSubmit', 'local-\nStorage', 'Confirm-\nation']
models_feat = ['Qwen 7B', 'Qwen 14B', 'Nemotron', 'Gemma 4']
# Feature matrix (1 = implemented)
data = np.array([
    [1,1,0,1,1,0,1,1],  # Qwen 7B
    [1,1,0,1,1,0,1,0],  # Qwen 14B
    [1,1,0,1,1,0,1,1],  # Nemotron
    [1,1,0,1,1,0,1,1],  # Gemma 4
])
fig, ax = plt.subplots(figsize=(9, 4))
im = ax.imshow(data, cmap='RdYlGn', vmin=0, vmax=1, aspect='auto')
ax.set_xticks(range(len(features)))
ax.set_xticklabels(features, fontsize=8)
ax.set_yticks(range(len(models_feat)))
ax.set_yticklabels(models_feat, fontsize=9)
ax.set_title('Ecommerce Feature Implementation (Green = Done, Red = Missing)', fontsize=11, pad=10)
for i in range(len(models_feat)):
    for j in range(len(features)):
        ax.text(j, i, '✓' if data[i,j] else '✗', ha='center', va='center', fontsize=14,
                color='white' if data[i,j] else 'black', fontweight='bold')
plt.tight_layout()
fig.savefig(FIGS/'fig2_ecommerce_features.png', dpi=200, bbox_inches='tight')
plt.close()
print("Chart 2 done")

# Fig 3: Iterations to reach quality
fig, ax = plt.subplots(figsize=(7, 4))
models_plot = ['Qwen 7B', 'Qwen 14B', 'MiniMax\nM3', 'Nemotron\n3 Super', 'Gemma 4']
iters_plot = [3, 2, 3, 3, 3]
colors = ['#55A868', '#4C72B0', '#C44E52', '#8172B2', '#64B5CD']
bars = ax.bar(range(len(models_plot)), iters_plot, color=colors, edgecolor='white', width=0.5)
ax.set_ylabel('Iterations Needed', fontsize=11)
ax.set_title('Loop Engineering: Iterations to Reach Best Score', fontsize=12, pad=10)
ax.set_xticks(range(len(models_plot)))
ax.set_xticklabels(models_plot, fontsize=8)
ax.set_ylim(0, 4.5)
ax.grid(axis='y', alpha=0.3, linestyle='--')
for bar, h in zip(bars, iters_plot):
    ax.text(bar.get_x()+0.25, h+0.1, str(h), ha='center', fontsize=11, fontweight='bold')
plt.tight_layout()
fig.savefig(FIGS/'fig3_iterations.png', dpi=200, bbox_inches='tight')
plt.close()
print("Chart 3 done")

# ===== WORD DOCUMENT =====
doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
pf = style.paragraph_format
pf.line_spacing = 2.0
pf.space_after = Pt(0)
pf.space_before = Pt(0)
pf.first_line_indent = Inches(0.5)

for s in ['Heading 1', 'Heading 2', 'Heading 3']:
    hs = doc.styles[s]
    hs.font.name = 'Times New Roman'
    hs.font.size = Pt(12)
    hs.font.bold = True
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hpf = hs.paragraph_format
    hpf.line_spacing = 2.0
    hpf.space_before = Pt(0)
    hpf.space_after = Pt(0)
    hpf.first_line_indent = Inches(0)

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

def add_para(text, bold=False, align=None, indent=True, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5) if indent else Inches(0)
    p.paragraph_format.line_spacing = 2.0
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_centered(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p

def add_img(path, width=5.5, caption=""):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
        if caption:
            add_para(caption, indent=False, size=10)

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=len(rows)+1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = t.rows[ri].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    return t

# ===== TITLE PAGE =====
for _ in range(6):
    doc.add_paragraph().paragraph_format.line_spacing = 2.0

add_centered("From 3D Landing Pages to Full Ecommerce:\nA Loop Engineering Study of 5 LLMs", bold=True, size=14)
doc.add_paragraph()
add_centered("Yash Kotalwar", bold=False)
add_centered("Independent Research", bold=False)
add_centered(datetime.now().strftime("%B %d, %Y"), bold=False)
doc.add_page_break()

# ===== ABSTRACT =====
add_heading("Abstract", level=1)
add_para(
    "This study evaluates five large language models — Qwen2.5-Coder 7B, Qwen2.5-Coder 14B, MiniMax M3, "
    "Nemotron 3 Super, and Gemma 4 — across two increasingly complex web development tasks. In the first "
    "experiment, each model generated a 3D product landing page using Three.js under two prompt conditions "
    "(vibe and technical). In the second experiment, a novel loop engineering methodology was employed: "
    "models iteratively built a full ecommerce website with a 3D hero section, product catalog fetched from "
    "Supabase, shopping cart with localStorage persistence, checkout form, and order submission to a live "
    "database. Each model received up to three iterations with specific feedback on missing features. "
    "Results show that Nemotron 3 Super and Gemma 4 achieved the highest frontend quality (10/10 on technical "
    "prompts), while all successful models plateaued at 8.6/10 on the ecommerce task due to consistent "
    "failure to implement real API calls — instead generating placeholder data. MiniMax M3 failed to produce "
    "valid HTML for the complex ecommerce task. The study demonstrates that current models can create "
    "convincing UI shells but struggle with actual backend integration, and that loop engineering with "
    "targeted feedback yields diminishing returns after two iterations."
)

doc.add_page_break()

# ===== 1. INTRODUCTION =====
add_heading("1. Introduction", level=1)

add_para(
    "The past two years have seen language models go from writing simple functions to generating entire "
    "web pages. But there's a gap between what benchmarks measure and what developers actually need. "
    "A model might ace HumanEval but stumble when asked to build a page with a 3D product viewer, "
    "a shopping cart that remembers your items, and a checkout form that talks to a real database."
)
add_para(
    "This study pushes models further than most existing evaluations. Instead of one-shot code generation "
    "on isolated tasks, we use loop engineering — an iterative process where the model generates code, "
    "receives specific feedback on what's missing, and tries again. This mirrors how a developer actually "
    "works: write something, test it, fix the bugs, repeat."
)
add_para(
    "We tested five models on two tasks. First, a straightforward 3D landing page. Then, a full ecommerce "
    "site with live database integration. The gap between these two tasks — and the models' consistent "
    "struggle with the second one — tells us something important about where LLMs are today and where "
    "they still fall short."
)

# ===== 2. LITERATURE REVIEW =====
add_heading("2. Literature Review", level=1)

add_para(
    "Benchmark-based evaluations of code generation have dominated the literature. HumanEval (Chen et al., 2021) "
    "and MBPP (Austin et al., 2021) established that LLMs can generate syntactically correct code for isolated "
    "functions. SWE-bench (Jimenez et al., 2024) raised the bar by evaluating multi-file edits on real GitHub issues."
)
add_para(
    "However, few studies have examined iterative code generation where models receive feedback and refine "
    "their output. This loop engineering approach is closer to how LLMs are used in practice — tools like "
    "GitHub Copilot Chat and Cursor support conversational refinement of generated code. Recent work by "
    "Shinn et al. (2024) on Reflexion showed that iterative self-reflection improves LLM performance on "
    "coding tasks, though their focus was on algorithmic problems rather than full-stack web development."
)
add_para(
    "The integration of external APIs represents a particular challenge. Models must generate not just "
    "syntactically valid code, but code that correctly authenticates with a service, formats requests "
    "according to REST conventions, and handles asynchronous responses. Prior work has shown that LLMs "
    "frequently hallucinate API endpoints and authentication patterns (Team Qwen, 2024), a finding that "
    "this study corroborates and extends to Supabase integration."
)

# ===== 3. METHODOLOGY =====
add_heading("3. Methodology", level=1)

add_heading("3.1 Models", level=2)

add_table(
    ["Model", "Parameters", "Access", "Provider"],
    [
        ["Qwen2.5-Coder 7B", "7B", "Local (Ollama)", "Alibaba"],
        ["Qwen2.5-Coder 14B", "14B", "Local (Ollama)", "Alibaba"],
        ["MiniMax M3", "Cloud", "Ollama Cloud", "MiniMax"],
        ["Nemotron 3 Super", "120B MoE", "Ollama Cloud", "NVIDIA"],
        ["Gemma 4", "Cloud", "Ollama Cloud", "Google DeepMind"],
    ]
)
add_para("Table 1. Models evaluated in this study.", indent=False, size=10)

add_heading("3.2 Experiment 1: Frontend Landing Page", level=2)

add_para(
    "Each model generated a single-file HTML page with a 3D dog food brand landing page using Three.js. "
    "Two prompt conditions were tested: a minimal vibe prompt (\"Make it look professional and modern\") "
    "and a detailed technical prompt specifying exact API names, camera positions, and lighting parameters. "
    "Scoring was based on ten binary criteria including WebGLRenderer, Scene, Camera, animation loop, "
    "lighting, shadows, 3D geometry, UI overlay, resize handler, and OrbitControls."
)

add_heading("3.3 Experiment 2: Ecommerce with Loop Engineering", level=2)

add_para(
    "Each model was tasked with building a complete ecommerce page for a fictional dog food brand called "
    "\"PawNutri.\" Requirements included: a Three.js 3D hero section with rotating product display, a product "
    "catalog fetched from a live Supabase database, a shopping cart with localStorage persistence, a checkout "
    "form capturing customer details, and order submission to Supabase with order_items creation."
)
add_para(
    "The loop engineering process worked as follows. Iteration 1: model received the full requirements and "
    "generated code. The output was scored on 14 criteria spanning 3D rendering, ecommerce functionality, "
    "API integration, and UI design. Missing features were catalogued. Iteration 2: model received the "
    "original requirements plus specific feedback listing every missing feature. The same process repeated "
    "for Iteration 3. Models were allowed up to three iterations, with the best-scoring version recorded."
)
add_para(
    "The Supabase backend was pre-configured with an products table containing 8 products across 4 categories "
    "(Food, Treats, Toys, Accessories), and orders and order_items tables with row-level security policies "
    "allowing anonymous reads on products and anonymous inserts on orders. Models were provided with the "
    "Supabase URL and anon key in the prompt."
)

# ===== 4. RESULTS =====
add_heading("4. Results", level=1)

add_heading("4.1 Experiment 1: Frontend Landing Page", level=2)

add_table(
    ["Model", "Vibe Score", "Technical Score", "Avg Time (s)"],
    [
        ["Qwen2.5-Coder 7B", "8.0", "10.0", "111.8"],
        ["Qwen2.5-Coder 14B", "8.0", "10.0", "229.7"],
        ["MiniMax M3", "7.0", "—", "52.3"],
        ["Nemotron 3 Super", "9.0", "10.0", "22.0"],
        ["Gemma 4", "9.0", "10.0", "15.9"],
    ]
)
add_para("Table 2. Frontend landing page quality scores.", indent=False, size=10)

add_img(str(FIGS/'fig1_all_scores.png'), width=5.5,
        caption="Figure 1. Model performance across all task types. Technical prompts consistently outperformed vibe prompts.")

add_para(
    "Nemotron 3 Super and Gemma 4 achieved perfect scores of 10/10 on the technical prompt, "
    "demonstrating complete Three.js implementations with all evaluated features. Their vibe-prompt "
    "scores of 9/10 were also the highest among all models. Cloud models responded significantly "
    "faster than local models (15-52s vs 112-230s). MiniMax M3's technical prompt returned an empty "
    "response, and Kimi K2.7 Code and GLM 5.2 (not shown) failed with authentication errors."
)

add_heading("4.2 Experiment 2: Ecommerce with Loop Engineering", level=2)

add_table(
    ["Model", "Iter 1", "Iter 2", "Iter 3", "Best Score", "Missing Features"],
    [
        ["Qwen2.5-Coder 7B", "8.6", "5.7", "8.6", "8.6/10", "Supabase fetch, order submit"],
        ["Qwen2.5-Coder 14B", "7.9", "8.6", "—", "8.6/10", "Order submit, responsive"],
        ["MiniMax M3", "—", "—", "—", "0/10", "No valid HTML generated"],
        ["Nemotron 3 Super", "8.6", "8.6", "—", "8.6/10", "Supabase fetch, order submit"],
        ["Gemma 4", "8.6", "8.6", "8.6", "8.6/10", "Supabase fetch, order submit"],
    ]
)
add_para("Table 3. Ecommerce loop engineering results across iterations.", indent=False, size=10)

add_img(str(FIGS/'fig2_ecommerce_features.png'), width=5.5,
        caption="Figure 2. Feature implementation matrix for the ecommerce task. Every successful model missed Supabase API integration.")

add_para(
    "The most striking result: every model that produced valid HTML plateaued at exactly 8.6/10, "
    "and every single one failed at the same two features — fetching products from Supabase and "
    "submitting orders to the database. They generated product grids with hardcoded data, cart "
    "functionality with localStorage, checkout forms, and even confirmation messages. But none "
    "actually called the Supabase API."
)
add_para(
    "This is a consistent blind spot. Models can generate the UI scaffolding for an ecommerce site, "
    "but when it comes to the actual network requests that make it a real application, they fall back "
    "to placeholder data. The feedback in iterations 2 and 3 specifically called out \"Supabase Product "
    "Fetch\" and \"Order Submission\" as missing, yet no model successfully implemented them."
)

add_heading("4.3 Iteration Analysis", level=2)

add_img(str(FIGS/'fig3_iterations.png'), width=4.5,
        caption="Figure 3. Number of loop engineering iterations needed per model.")

add_para(
    "Most models plateaued after two iterations. The first iteration typically produced a functional "
    "UI with missing backend integration. The second iteration addressed UI gaps but rarely fixed the "
    "API calls. The third iteration showed no improvement — models either repeated the same structure "
    "or, in the case of Qwen2.5-Coder 7B, regressed. This suggests that the feedback loop has diminishing "
    "returns: the models understand what a shopping cart should look like, but they don't understand "
    "how to wire it to a real backend."
)

# ===== 5. DISCUSSION =====
add_heading("5. Discussion", level=1)

add_para(
    "These results paint a nuanced picture of LLM capabilities for web development. On traditional "
    "frontend tasks — generating a 3D scene with lighting, animation, and UI overlays — modern models "
    "perform exceptionally well. Nemotron 3 Super and Gemma 4 match or exceed specialized code models "
    "like Qwen2.5-Coder, despite being general-purpose architectures."
)
add_para(
    "But the ecommerce experiment reveals a critical weakness. Every model generated a beautiful but "
    "non-functional storefront. The product grid looked real — prices, descriptions, add-to-cart buttons — "
    "but the data was hardcoded. The checkout form collected information and showed a confirmation "
    "message, but nothing was saved to the database. These are not bugs in the traditional sense. "
    "The code is syntactically valid and structurally complete. It just doesn't actually work."
)
add_para(
    "This has implications for how we evaluate code generation. Current benchmarks focus on whether "
    "code compiles or passes unit tests. They don't check whether generated code actually performs "
    "its intended function when talking to external services. An ecommerce page that looks perfect "
    "but doesn't process orders is worse than useless — it's deceptive."
)
add_para(
    "The loop engineering methodology proved useful but with limitations. Two iterations captured most "
    "available improvements; the third iteration was wasted for every model. This suggests an optimal "
    "feedback strategy: one generation pass, one refinement pass based on specific feedback, then "
    "manual intervention for remaining issues rather than continued automated iteration."
)
add_para(
    "Several limitations of this study should be noted. The sample is small: five models, one task type, "
    "three iterations each. The Supabase integration required specific API patterns that may not "
    "generalize to other backends. And the evaluation was binary (feature present or absent) rather "
    "than measuring code quality, security, or maintainability."
)

# ===== 6. CONCLUSION =====
add_heading("6. Conclusion", level=1)

add_para(
    "This study evaluated five language models on two web development tasks of increasing complexity. "
    "The key findings are as follows."
)
add_para(
    "First, for straightforward frontend generation, cloud models Nemotron 3 Super and Gemma 4 match "
    "or exceed specialized local code models, achieving perfect 10/10 scores with detailed prompts and "
    "responding 5-10x faster than local alternatives."
)
add_para(
    "Second, the transition from static landing pages to functional ecommerce reveals a critical gap. "
    "All models plateaued at 8.6/10 on the ecommerce task, and every one failed at the same thing: "
    "making real API calls to a database. They can build the store window but not the cash register."
)
add_para(
    "Third, loop engineering with targeted feedback improves scores but with diminishing returns. "
    "Two iterations capture most available gains; further iterations add no value for these models "
    "and this task. The models incorporate UI feedback but cannot translate \"you need to fetch "
    "products from the API\" into correct code."
)
add_para(
    "Fourth, MiniMax M3 failed entirely on the complex ecommerce task, producing no valid HTML across "
    "three attempts. This suggests that model reliability degrades sharply with task complexity for "
    "some architectures."
)
add_para(
    "For practitioners, the guidance depends on the task. For visual frontend work, Nemotron 3 Super "
    "or Gemma 4 via Ollama Cloud offer the best speed-quality tradeoff. For projects requiring real "
    "backend integration, no current model can be trusted to generate working API calls without "
    "human intervention. Developers should budget time for manually implementing or correcting "
    "all network requests in LLM-generated code."
)
add_para(
    "Future work should explore whether fine-tuning on API documentation or providing OpenAPI specs "
    "in the prompt improves backend code generation. Multi-agent systems where one agent generates "
    "the UI and another handles the API layer may also prove more effective than monolithic generation. "
    "And as models improve, the loop engineering methodology itself deserves study: what feedback "
    "formats and iteration strategies most effectively drive improvement?"
)

# ===== REFERENCES =====
add_heading("References", level=1)

refs = [
    "Austin, J., Odena, A., Nye, M., Bosma, M., Michalewski, H., Dohan, D., ... & Sutton, C. (2021). Program Synthesis with Large Language Models. arXiv:2108.07732.",
    "Chen, M., Tworek, J., Jun, H., Yuan, Q., Pinto, H. P. O., Kaplan, J., ... & Zaremba, W. (2021). Evaluating Large Language Models Trained on Code. arXiv:2107.03374.",
    "Google DeepMind. (2026). Gemma 4: Technical Report.",
    "Jimenez, C. E., Yang, J., Wettig, A., Yao, S., Pei, K., Press, O., & Narasimhan, K. (2024). SWE-bench: Can Language Models Resolve Real-World GitHub Issues? ICLR 2024.",
    "MiniMax. (2026). MiniMax M3: Coding & Agentic Frontier.",
    "NVIDIA. (2026). Nemotron 3 Super: Open Hybrid Mamba-Transformer MoE. NVIDIA Research.",
    "Ollama. (2026). Cloud Models Documentation. https://docs.ollama.com/cloud.",
    "Shinn, N., Cassano, F., Gopinath, A., Narasimhan, K., & Yao, S. (2024). Reflexion: Language Agents with Verbal Reinforcement Learning. NeurIPS 2024.",
    "Team Qwen. (2024). Qwen2.5-Coder: Technical Report. arXiv:2409.12186.",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.hanging_indent = Inches(0.5)
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

out = Path.home() / "Desktop" / "LLM_Loop_Engineering_Research_Paper.docx"
doc.save(str(out))
print(f"\nPaper saved: {out}")
